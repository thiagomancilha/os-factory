#!/usr/bin/env python3
"""Tests for tools/build_os_docx.py and tools/os_registry.py.

Uses fully synthetic fixtures (generic demand id, generic prose, a
generated placeholder PNG, isolated temp config/registry files) — no
real client, project or contract data. Runs with the stdlib `unittest`
runner so it has no extra dependency beyond what `requirements.txt`
already declares:

    python3 -m unittest tests.test_build_os_docx -v
    (run from the repository root)
"""
from __future__ import annotations

import json
import shutil
import sys
import tempfile
import unittest
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT / "tools"))

from docx import Document  # noqa: E402
from docx.shared import Pt  # noqa: E402
from PIL import Image  # noqa: E402

import build_os_docx as bod  # noqa: E402
import os_commercial  # noqa: E402
import os_registry  # noqa: E402

# A markdown fixture that is deliberately generic: fictitious demand id,
# fictitious section content, no company/person/contract names.
FIXTURE_MARKDOWN = """# Documento de Especificação Funcional

## Identificação

- Código / Identificador da demanda: 0000-00-00-demo
- Nome da demanda: Demanda de exemplo genérica
- Projeto: Projeto Exemplo
- Solicitante: Cliente Genérico Ltda
- Responsável: Analista de Exemplo
- Data: 01/01/2026

## Controle de versão

| Versão | Data | Autor | Descrição |
| ------ | ---- | ----- | --------- |
| 1.0 | 01/01/2026 | pipeline-de-teste | Primeira versão de teste. |

## 1. DESCRIÇÃO

Texto de descrição genérico para fins de teste automatizado.

## 3. ESCOPO

`[FIGURA: Diagrama de exemplo | fonte: diagrama.png]`

### 3.1 Regras Funcionais

- Primeira regra de exemplo, com `identificador_tecnico` embutido.
- Segunda regra de exemplo. `[DISCOVERY_ITEM D1: item de exemplo — a descobrir na etapa de teste.]`

**Condição:** condição de exemplo.
**Comportamento esperado:** comportamento de exemplo. `[OPEN_QUESTION: pendência de exemplo?]`

`[ARCHITECTURAL_CONFLICT: divergência de exemplo entre duas fontes fictícias.]`

| Método / rota | Finalidade |
| --- | --- |
| `GET /api/v1/exemplo` | Endpoint de exemplo. |
| `GET /health` | Health check de exemplo. |

## 6. ESFORÇO

| Descrição | Esforço |
| --- | ---: |
| Total | 10 horas |

## 7. CONDIÇÕES COMERCIAIS

- Horas contratadas: 10 horas
- Valor total da OS: R$ 1.000,00
- Forma / condição de pagamento: 50% no início, 50% na entrega

## 8. ACEITE

- Método de aceite: E-mail.

O aceite desta Ordem de Serviço poderá ser formalizado por e-mail.
"""

FIXTURE_MARKDOWN_MISSING_FIGURE = FIXTURE_MARKDOWN.replace(
    "fonte: diagrama.png", "fonte: arquivo-que-nao-existe.png"
)

FIXTURE_MARKDOWN_NO_VALOR_PAGAMENTO = FIXTURE_MARKDOWN.replace(
    "- Valor total da OS: R$ 1.000,00\n- Forma / condição de pagamento: 50% no início, 50% na entrega",
    "- Valor total da OS: `[OPEN_QUESTION: valor ainda não informado.]`\n"
    "- Forma / condição de pagamento: `[OPEN_QUESTION: forma de pagamento ainda não informada.]`",
)

# Same commercial figures as FIXTURE_MARKDOWN, but with NO specific forma de
# pagamento given (only horas + valor) — used to test that, up to the
# policy's default threshold, the Tria standard 50/50 policy applies when no
# demand-specific condition is present.
FIXTURE_MARKDOWN_NO_PAGAMENTO = FIXTURE_MARKDOWN.replace(
    "- Forma / condição de pagamento: 50% no início, 50% na entrega",
    "- Forma / condição de pagamento: `[OPEN_QUESTION: forma de pagamento ainda não informada.]`",
)

# Same as FIXTURE_MARKDOWN_NO_PAGAMENTO but with an effort above the policy's
# default threshold (80h in the real Tria policy; the test policy below uses
# the same limit) — used to test that no installment plan is invented above
# the threshold; PAYMENT_BY_MILESTONES_RECOMMENDED must surface instead.
FIXTURE_MARKDOWN_HIGH_EFFORT_NO_PAGAMENTO = FIXTURE_MARKDOWN_NO_PAGAMENTO.replace(
    "| Total | 10 horas |", "| Total | 120 horas |"
).replace(
    "- Horas contratadas: 10 horas", "- Horas contratadas: 120 horas"
)

# A synthetic commercial policy fixture — isolated from the real
# config/os-commercial-policy.json so these tests never depend on (or risk
# being broken by) changes to Tria's actual institutional policy.
TEST_COMMERCIAL_POLICY = {
    "limite_horas_pagamento_padrao": 80,
    "parcelas_padrao": [
        {"marco": "no aceite da Ordem de Serviço", "percentual": 50},
        {"marco": "na entrega final", "percentual": 50},
    ],
    "marcos_sugeridos_referencia": [
        "aceite / kick-off",
        "conclusão de Discovery",
        "entrega final",
    ],
}

FORBIDDEN_SUBSTRINGS = [
    "SAGI",
    "GCONV",
    "Fundação PÁTRIA",
    "Quanthum",
    "Entendimento Técnico",
]

TEST_INSTITUTIONAL_CONFIG = {
    "contratante_padrao": "Empresa Genérica de Teste",
    "executor_padrao": "Executor de Teste",
    "autor_padrao": "Executor de Teste",
    "formato_codigo_os": "OS-AAAA-NNNN",
    "metodo_padrao_aceite": "E-mail",
}


def _build_minimal_template(path: Path, with_cover: bool = False) -> None:
    """Builds a tiny, fully generic DOCX template for test isolation.

    Deliberately independent from `04-templates/docx/os-padrao.docx` so
    these tests never depend on (or risk mutating) the real official
    template, and stay fast/self-contained. When `with_cover` is set, a
    minimal cover page (title + metadata table matching
    `bod.COVER_LABELS`) is included, mirroring the shape of the official
    template's cover — used to test cover-filling in isolation.
    """
    doc = Document()
    for name in ("Normal", "Title", "Heading 1", "Heading 2", "List Bullet", "Caption"):
        style = doc.styles[name]
        style.font.size = Pt(11 if name in ("Normal", "List Bullet") else 13)

    if with_cover:
        title_p = doc.add_paragraph()
        title_p.add_run("ORDEM DE SERVIÇO")
        doc.add_paragraph("Especificação Funcional e Condições de Execução")
        table = doc.add_table(rows=len(bod.COVER_LABELS), cols=2)
        for row, label in zip(table.rows, bod.COVER_LABELS):
            row.cells[0].text = f"{label.capitalize()}:"
            row.cells[1].text = ""
        doc.add_page_break()

    placeholder = doc.add_paragraph("[template oficial de teste]")
    doc.save(str(path))


class BuildOsDocxTestCase(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp(prefix="osfactory_test_"))
        self.template_path = self.tmp / "template.docx"
        _build_minimal_template(self.template_path, with_cover=False)

        self.cover_template_path = self.tmp / "template_cover.docx"
        _build_minimal_template(self.cover_template_path, with_cover=True)

        self.figures_dir = self.tmp / "00-inbox" / "0000-00-00-demo"
        self.figures_dir.mkdir(parents=True)
        img = Image.new("RGB", (800, 400), color=(210, 210, 210))
        img.save(self.figures_dir / "diagrama.png")

        self.markdown_path = self.tmp / "OS-0000-00-00-demo.md"
        self.markdown_path.write_text(FIXTURE_MARKDOWN, encoding="utf-8")

        self.output_path = self.tmp / "out" / "OS-0000-00-00-demo.docx"

        self.config_path = self.tmp / "os-factory.json"
        self.config_path.write_text(json.dumps(TEST_INSTITUTIONAL_CONFIG), encoding="utf-8")

        self.registry_path = self.tmp / "os-registry.json"

        self.commercial_policy_path = self.tmp / "os-commercial-policy.json"
        self.commercial_policy_path.write_text(json.dumps(TEST_COMMERCIAL_POLICY), encoding="utf-8")
        self.commercial_registry_path = self.tmp / "os-commercial.json"

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _build(self, markdown_path=None, template_path=None, output_path=None,
               demand="0000-00-00-demo", config_path=None, registry_path=None,
               validation_status=None, **extra):
        kwargs = dict(
            demand=demand,
            markdown_path=markdown_path or self.markdown_path,
            template_path=template_path or self.template_path,
            output_path=output_path or self.output_path,
            figures_dir=self.figures_dir,
            validation_status=validation_status,
            config_path=config_path or self.config_path,
            registry_path=registry_path or self.registry_path,
            explicit_os_code=None,
            explicit_year="2026",
            approved=False,
            approved_by=None,
            approved_date=None,
            acceptance_reference=None,
            commercial_policy_path=self.commercial_policy_path,
            commercial_registry_path=self.commercial_registry_path,
        )
        kwargs.update(extra)
        return bod.build(**kwargs)

    # 1. template original não é modificado
    def test_template_not_modified(self):
        before = self.template_path.read_bytes()
        self._build()
        after = self.template_path.read_bytes()
        self.assertEqual(before, after)

    # 2. DOCX é criado
    def test_docx_is_created(self):
        self._build()
        self.assertTrue(self.output_path.is_file())

    # 3. headings são materializados
    def test_headings_materialized(self):
        self._build()
        doc = Document(str(self.output_path))
        styles_seen = {p.style.name for p in doc.paragraphs}
        self.assertIn("Title", styles_seen)
        self.assertIn("Heading 1", styles_seen)
        self.assertIn("Heading 2", styles_seen)

    # 4. texto essencial aparece
    def test_essential_text_present(self):
        self._build()
        doc = Document(str(self.output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Texto de descrição genérico para fins de teste automatizado.", full_text)

    # 5. bullets funcionam
    def test_bullets_render(self):
        self._build()
        doc = Document(str(self.output_path))
        bullets = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
        self.assertTrue(any("Primeira regra de exemplo" in b for b in bullets))

    # 6. tabela funciona
    def test_table_renders(self):
        self._build()
        doc = Document(str(self.output_path))
        endpoints_tables = [t for t in doc.tables if t.cell(0, 0).text == "Método / rota"]
        self.assertEqual(len(endpoints_tables), 1)
        self.assertEqual(endpoints_tables[0].cell(1, 0).text, "GET /api/v1/exemplo")

    # 7. subseções dinâmicas funcionam
    def test_dynamic_subsection_renders(self):
        self._build()
        doc = Document(str(self.output_path))
        h2_texts = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
        self.assertIn("3.1 Regras Funcionais", h2_texts)

    # 8. figura válida é inserida
    def test_valid_figure_inserted(self):
        self._build()
        doc = Document(str(self.output_path))
        self.assertEqual(len(doc.inline_shapes), 1)
        captions = [p.text for p in doc.paragraphs if p.style.name == "Caption"]
        self.assertTrue(any("Diagrama de exemplo" in c for c in captions))

    # 9. figura inexistente gera erro controlado
    def test_missing_figure_raises_materialization_error(self):
        bad_markdown = self.tmp / "OS-bad.md"
        bad_markdown.write_text(FIXTURE_MARKDOWN_MISSING_FIGURE, encoding="utf-8")
        with self.assertRaises(bod.MaterializationError):
            self._build(markdown_path=bad_markdown)
        self.assertFalse(self.output_path.exists())

    # 10. marcadores de incerteza/conflito são preservados
    def test_uncertainty_and_conflict_markers_preserved(self):
        self._build()
        doc = Document(str(self.output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("[DISCOVERY_ITEM D1: item de exemplo — a descobrir na etapa de teste.]", full_text)
        self.assertIn("[OPEN_QUESTION: pendência de exemplo?]", full_text)
        self.assertIn("[ARCHITECTURAL_CONFLICT: divergência de exemplo entre duas fontes fictícias.]", full_text)
        marker_runs = [
            r.text
            for p in doc.paragraphs
            for r in p.runs
            if r.italic and r.font.name == "Consolas"
        ]
        self.assertTrue(any("DISCOVERY_ITEM" in m for m in marker_runs))

    # 11. nenhum nome de cliente está hardcoded
    def test_no_hardcoded_client_names_in_tool_source(self):
        source = (REPO_ROOT / "tools" / "build_os_docx.py").read_text(encoding="utf-8")
        for forbidden in FORBIDDEN_SUBSTRINGS:
            self.assertNotIn(forbidden, source)

    # Functional gate: BLOCKED must never produce a DOCX.
    def test_functional_blocked_refuses_to_build(self):
        with self.assertRaises(bod.FunctionalGateError):
            self._build(validation_status="BLOCKED")
        self.assertFalse(self.output_path.exists())

    # Determinism: same inputs -> byte-identical output.
    def test_deterministic_output(self):
        self._build()
        first = self.output_path.read_bytes()
        second_output = self.tmp / "out2" / "OS-0000-00-00-demo.docx"
        second_registry = self.tmp / "os-registry-2.json"
        self._build(output_path=second_output, registry_path=second_registry)
        second = second_output.read_bytes()
        self.assertEqual(first, second)

    # --- Código da OS / registry --------------------------------------

    def test_same_demand_reuses_same_code(self):
        r1 = self._build()
        second_output = self.tmp / "out2" / "OS.docx"
        r2 = self._build(output_path=second_output)
        self.assertEqual(r1["os_code"], r2["os_code"])

    def test_next_demand_gets_next_sequential(self):
        r1 = self._build()
        md2 = self.tmp / "OS-outro.md"
        md2.write_text(FIXTURE_MARKDOWN.replace("0000-00-00-demo", "0000-00-00-outro"), encoding="utf-8")
        r2 = self._build(demand="0000-00-00-outro", markdown_path=md2, output_path=self.tmp / "out2" / "OS.docx")
        year1, seq1 = r1["os_code"].split("-")[1:]
        year2, seq2 = r2["os_code"].split("-")[1:]
        self.assertEqual(year1, year2)
        self.assertEqual(int(seq2), int(seq1) + 1)

    def test_year_change_resets_sequence(self):
        code_2026 = os_registry.resolve_os_code("demanda-a", "2026", registry_path=self.registry_path)
        code_2027 = os_registry.resolve_os_code("demanda-b", "2027", registry_path=self.registry_path)
        self.assertTrue(code_2026.startswith("OS-2026-"))
        self.assertTrue(code_2027.startswith("OS-2027-0001"))

    def test_registry_never_reuses_code_from_another_demand(self):
        code = os_registry.resolve_os_code("demanda-a", "2026", registry_path=self.registry_path)
        with self.assertRaises(os_registry.RegistryError):
            os_registry.resolve_os_code("demanda-b", "2026", explicit_code=code, registry_path=self.registry_path)

    # --- Defaults institucionais ---------------------------------------

    def test_executor_padrao_is_thiago_mancilha(self):
        real_config = REPO_ROOT / "config" / "os-factory.json"
        config = json.loads(real_config.read_text(encoding="utf-8"))
        self.assertEqual(config["executor_padrao"], "Thiago Mancilha")

    def test_contratante_padrao_is_tria(self):
        real_config = REPO_ROOT / "config" / "os-factory.json"
        config = json.loads(real_config.read_text(encoding="utf-8"))
        self.assertEqual(config["contratante_padrao"], "Tria")

    # --- Capa / identidade -----------------------------------------------

    def test_cover_contains_ordem_de_servico(self):
        self._build(template_path=self.cover_template_path)
        doc = Document(str(self.output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("ORDEM DE SERVIÇO", full_text)

    def test_cover_table_filled_with_os_code(self):
        result = self._build(template_path=self.cover_template_path)
        doc = Document(str(self.output_path))
        cover_table = doc.tables[0]
        values = {row.cells[0].text.strip().rstrip(":").lower(): row.cells[1].text for row in cover_table.rows}
        self.assertEqual(values.get("código da os"), result["os_code"])
        self.assertEqual(values.get("executor"), TEST_INSTITUTIONAL_CONFIG["executor_padrao"])
        self.assertEqual(values.get("contratante"), TEST_INSTITUTIONAL_CONFIG["contratante_padrao"])

    # --- Condições Comerciais / Aceite ------------------------------------

    def test_condicoes_comerciais_section_exists(self):
        self._build()
        doc = Document(str(self.output_path))
        headings = [p.text for p in doc.paragraphs if p.style.name in ("Heading 1", "Title")]
        self.assertTrue(any("CONDIÇÕES COMERCIAIS" in h for h in headings))

    def test_aceite_section_exists(self):
        self._build()
        doc = Document(str(self.output_path))
        headings = [p.text for p in doc.paragraphs if p.style.name in ("Heading 1", "Title")]
        self.assertTrue(any(h.strip().upper().endswith("ACEITE") for h in headings))

    def test_metodo_aceite_padrao_is_email(self):
        result = self._build()
        self.assertEqual(result["metodo_aceite"], "E-mail")

    # --- Readiness para aceite -------------------------------------------

    def test_missing_valor_results_in_not_ready(self):
        md = self.tmp / "OS-sem-valor.md"
        md.write_text(FIXTURE_MARKDOWN_NO_VALOR_PAGAMENTO, encoding="utf-8")
        result = self._build(markdown_path=md, validation_status="PASS_WITH_WARNINGS")
        self.assertEqual(result["acceptance_readiness"], "NOT_READY_FOR_ACCEPTANCE")
        self.assertIn("Valor da OS", result["missing"])

    def test_missing_pagamento_results_in_not_ready(self):
        md = self.tmp / "OS-sem-pagamento.md"
        md.write_text(FIXTURE_MARKDOWN_NO_VALOR_PAGAMENTO, encoding="utf-8")
        result = self._build(markdown_path=md, validation_status="PASS_WITH_WARNINGS")
        self.assertIn("Forma de pagamento", result["missing"])

    def test_complete_commercial_fields_result_in_ready(self):
        result = self._build(validation_status="PASS_WITH_WARNINGS")
        self.assertEqual(result["acceptance_readiness"], "READY_FOR_ACCEPTANCE")
        self.assertEqual(result["missing"], [])

    # --- Hierarquia da condição comercial (ver os-rules.md, OS-COMMERCIAL-001) ---

    def test_specific_condition_in_insumos_prevails_over_default_policy(self):
        # FIXTURE_MARKDOWN already carries a specific forma de pagamento
        # ("50% no início, 50% na entrega") distinct from the Tria default
        # policy wording — it must win over the policy default, unmodified.
        result = self._build(validation_status="PASS_WITH_WARNINGS")
        pagamento = result["commercial"]["forma/condição de pagamento"]
        self.assertEqual(pagamento["source"], "insumos")
        self.assertEqual(pagamento["value"], "50% no início, 50% na entrega")

    def test_explicit_confirmation_prevails_over_default_policy(self):
        md = self.tmp / "OS-sem-pagamento.md"
        md.write_text(FIXTURE_MARKDOWN_NO_PAGAMENTO, encoding="utf-8")
        result = self._build(
            markdown_path=md,
            validation_status="PASS_WITH_WARNINGS",
            explicit_pagamento_splits=[
                {"marco": "na assinatura", "percentual": 30},
                {"marco": "na entrega", "percentual": 70},
            ],
        )
        pagamento = result["commercial"]["forma/condição de pagamento"]
        self.assertEqual(pagamento["source"], "usuario_confirmado")
        self.assertIn("30%", pagamento["value"])
        self.assertIn("70%", pagamento["value"])
        self.assertNotIn("aceite da Ordem de Serviço", pagamento["value"])

    def test_explicit_confirmation_supersedes_insumos_without_conflict(self):
        # FIXTURE_MARKDOWN's insumos already state "R$ 1.000,00"; an explicit,
        # later confirmation of a different value must simply prevail — this
        # is supersession, not a CONFLICT (hierarchy already ranks explicit
        # confirmation above insumos).
        result = self._build(
            validation_status="PASS_WITH_WARNINGS",
            explicit_valor="R$ 8.000,00",
        )
        valor = result["commercial"]["valor total da os"]
        self.assertEqual(valor["source"], "usuario_confirmado")
        self.assertEqual(valor["value"], "R$ 8.000,00")

    def test_default_5050_applies_up_to_policy_limit_when_unspecified(self):
        md = self.tmp / "OS-sem-pagamento.md"
        md.write_text(FIXTURE_MARKDOWN_NO_PAGAMENTO, encoding="utf-8")
        result = self._build(markdown_path=md, validation_status="PASS_WITH_WARNINGS")
        pagamento = result["commercial"]["forma/condição de pagamento"]
        self.assertEqual(pagamento["source"], "politica_padrao_tria")
        self.assertIn("50%", pagamento["value"])
        self.assertIn("R$ 500,00", pagamento["value"])

    def test_above_policy_limit_does_not_invent_installments(self):
        md = self.tmp / "OS-esforco-alto.md"
        md.write_text(FIXTURE_MARKDOWN_HIGH_EFFORT_NO_PAGAMENTO, encoding="utf-8")
        result = self._build(markdown_path=md, validation_status="PASS_WITH_WARNINGS")
        pagamento = result["commercial"]["forma/condição de pagamento"]
        self.assertEqual(pagamento["source"], os_commercial.PAYMENT_BY_MILESTONES_RECOMMENDED)
        self.assertIn("OPEN_QUESTION", pagamento["value"])
        self.assertNotIn("R$", pagamento["value"])
        self.assertEqual(result["acceptance_readiness"], "NOT_READY_FOR_ACCEPTANCE")
        self.assertTrue(any("PAYMENT_BY_MILESTONES_RECOMMENDED" in m for m in result["missing"]))

    def test_installments_reconcile_exactly_with_total_value(self):
        installments = os_commercial.compute_installments(
            "R$ 8.000,00",
            [
                {"marco": "no aceite da Ordem de Serviço", "percentual": 50},
                {"marco": "na entrega final", "percentual": 50},
            ],
        )
        total = sum(os_commercial.parse_brl(i["valor"]) for i in installments)
        self.assertEqual(total, os_commercial.parse_brl("R$ 8.000,00"))

    def test_external_docx_never_shows_superseded_commercial_value(self):
        # FIXTURE_MARKDOWN's insumos state "R$ 1.000,00"; an explicit later
        # confirmation of a different value must fully replace it in the
        # DOCX handed to the client — the superseded value must not leak
        # anywhere into the external document (see os-rules.md,
        # OS-COMMERCIAL-003). Internal traceability of the supersession
        # belongs only to internal artifacts (validation.md, runtime
        # registry), never to the OS itself.
        self._build(validation_status="PASS_WITH_WARNINGS", explicit_valor="R$ 8.000,00")
        doc = Document(str(self.output_path))
        full_text = "\n".join(
            [p.text for p in doc.paragraphs]
            + [cell.text for t in doc.tables for row in t.rows for cell in row.cells]
        )
        self.assertNotIn("1.000,00", full_text)
        self.assertIn("8.000,00", full_text)

    def test_new_explicit_commercial_confirmation_keeps_same_os_code(self):
        r1 = self._build(validation_status="PASS_WITH_WARNINGS")
        second_output = self.tmp / "out2" / "OS.docx"
        r2 = self._build(
            output_path=second_output,
            validation_status="PASS_WITH_WARNINGS",
            explicit_horas="80 horas",
            explicit_valor="R$ 8.000,00",
            explicit_pagamento_splits=[
                {"marco": "no aceite da Ordem de Serviço", "percentual": 50},
                {"marco": "na entrega final", "percentual": 50},
            ],
        )
        self.assertEqual(r1["os_code"], r2["os_code"])

    # --- Aprovação nunca é inventada ---------------------------------------

    def test_approval_never_inferred_by_default(self):
        result = self._build(validation_status="PASS_WITH_WARNINGS")
        self.assertNotEqual(result["document_status"], "Aprovada")

    def test_approved_requires_approved_by(self):
        with self.assertRaises(bod.MaterializationError):
            self._build(approved=True, approved_by=None)

    def test_approved_explicit_sets_aprovada(self):
        result = self._build(approved=True, approved_by="Cliente Genérico Ltda")
        self.assertEqual(result["document_status"], "Aprovada")


if __name__ == "__main__":
    unittest.main()
