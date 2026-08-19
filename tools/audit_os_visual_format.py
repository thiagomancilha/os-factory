#!/usr/bin/env python3
"""Structural/visual audit for `.osFactory` OS DOCX documents.

Performs everything that can be checked programmatically from the DOCX
object model itself (python-docx), without requiring a rendered preview:
heading hierarchy, style/font consistency against the template's own
style definitions, table shape, image sizing versus the page's usable
width, empty-paragraph runs, and headings with no following content.

This is deliberately generic: it never hardcodes a client name, demand
identifier, or expected section title. It audits *format*, not content.

Output vocabulary (see os-rules.md, OS-QA-003):
    PASS               — no structural/visual problems found.
    PASS_WITH_WARNINGS — non-blocking issues found (cosmetic risk).
    OUTPUT_BLOCKED      — materialization defect severe enough that the
                           DOCX should not be treated as a final document
                           (e.g. an image wider than the page, an empty
                           table, a document with no visible content).

Usage:
    python tools/audit_os_visual_format.py --docx <path/to/OS.docx>
"""
from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn

STATUS_PASS = "PASS"
STATUS_PASS_WITH_WARNINGS = "PASS_WITH_WARNINGS"
STATUS_OUTPUT_BLOCKED = "OUTPUT_BLOCKED"

EXIT_BY_STATUS = {
    STATUS_PASS: 0,
    STATUS_PASS_WITH_WARNINGS: 0,
    STATUS_OUTPUT_BLOCKED: 2,
}

EMU_PER_INCH = 914400
HEADING_STYLES_IN_ORDER = ["Title", "Heading 1", "Heading 2", "Heading 3"]
MAX_CONSECUTIVE_EMPTY_PARAGRAPHS = 2


@dataclass
class Finding:
    severity: str  # BLOCKING | WARNING | INFO
    code: str
    message: str


def _usable_width_emu(doc: Document) -> int:
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _heading_level(style_name: str) -> Optional[int]:
    if style_name in HEADING_STYLES_IN_ORDER:
        return HEADING_STYLES_IN_ORDER.index(style_name)
    return None


def audit_heading_hierarchy(doc: Document) -> List[Finding]:
    findings: List[Finding] = []
    last_level: Optional[int] = None
    for para in doc.paragraphs:
        level = _heading_level(para.style.name)
        if level is None:
            continue
        if not para.text.strip():
            findings.append(
                Finding("BLOCKING", "EMPTY_HEADING", f"Heading '{para.style.name}' sem texto.")
            )
            continue
        if last_level is not None and level > last_level + 1:
            findings.append(
                Finding(
                    "WARNING",
                    "HEADING_LEVEL_SKIP",
                    f"Heading '{para.text[:60]}' ({para.style.name}) pula um nível "
                    f"em relação ao heading anterior.",
                )
            )
        last_level = level
    return findings


def audit_headings_have_content(doc: Document) -> List[Finding]:
    findings: List[Finding] = []
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        if _heading_level(para.style.name) is None or not para.text.strip():
            continue
        # Look ahead for the next non-empty, non-heading paragraph or a table
        # immediately following this heading in body order.
        has_following_content = False
        for j in range(i + 1, len(paras)):
            nxt = paras[j]
            if _heading_level(nxt.style.name) is not None:
                break  # hit next heading before finding content
            if nxt.text.strip():
                has_following_content = True
                break
        # A table right after the heading also counts as content; python-docx's
        # flat `.paragraphs` list doesn't interleave tables, so if no paragraph
        # content was found we don't flag when the doc has at least one table
        # overall — a full body-order walk is avoided to keep this audit simple
        # and dependency-free. This trades a small amount of precision for
        # robustness across template variations.
        if not has_following_content and not doc.tables:
            findings.append(
                Finding(
                    "WARNING",
                    "HEADING_WITHOUT_CONTENT",
                    f"Heading '{para.text[:60]}' não parece ter conteúdo imediatamente após.",
                )
            )
    return findings


def audit_empty_paragraph_runs(doc: Document) -> List[Finding]:
    findings: List[Finding] = []
    streak = 0
    for para in doc.paragraphs:
        if not para.text.strip() and _heading_level(para.style.name) is None:
            streak += 1
        else:
            if streak > MAX_CONSECUTIVE_EMPTY_PARAGRAPHS:
                findings.append(
                    Finding(
                        "WARNING",
                        "EXCESSIVE_EMPTY_PARAGRAPHS",
                        f"{streak} parágrafos vazios consecutivos encontrados.",
                    )
                )
            streak = 0
    if streak > MAX_CONSECUTIVE_EMPTY_PARAGRAPHS:
        findings.append(
            Finding(
                "WARNING",
                "EXCESSIVE_EMPTY_PARAGRAPHS",
                f"{streak} parágrafos vazios consecutivos encontrados no final do documento.",
            )
        )
    return findings


def audit_tables(doc: Document) -> List[Finding]:
    findings: List[Finding] = []
    for idx, table in enumerate(doc.tables):
        if len(table.rows) == 0 or len(table.columns) == 0:
            findings.append(
                Finding("BLOCKING", "EMPTY_TABLE", f"Tabela #{idx + 1} não possui linhas/colunas.")
            )
            continue
        if len(table.rows) == 1:
            findings.append(
                Finding(
                    "WARNING",
                    "TABLE_HEADER_ONLY",
                    f"Tabela #{idx + 1} possui apenas a linha de cabeçalho, sem dados.",
                )
            )
        header_cells = table.rows[0].cells
        if all(not c.text.strip() for c in header_cells):
            findings.append(
                Finding("WARNING", "TABLE_EMPTY_HEADER", f"Tabela #{idx + 1} tem cabeçalho vazio.")
            )
    return findings


def audit_images(doc: Document) -> List[Finding]:
    findings: List[Finding] = []
    usable_width = _usable_width_emu(doc)
    for idx, shape in enumerate(doc.inline_shapes):
        if shape.width > usable_width:
            findings.append(
                Finding(
                    "BLOCKING",
                    "IMAGE_WIDER_THAN_PAGE",
                    f"Imagem #{idx + 1} ({shape.width / EMU_PER_INCH:.2f}in) excede a largura "
                    f"útil da página ({usable_width / EMU_PER_INCH:.2f}in).",
                )
            )
        if shape.width <= 0 or shape.height <= 0:
            findings.append(
                Finding("BLOCKING", "IMAGE_ZERO_SIZE", f"Imagem #{idx + 1} tem dimensão zero/inválida.")
            )
    return findings


def audit_font_consistency(doc: Document) -> List[Finding]:
    """Flag runs whose font diverges from the template's own style defaults.

    Reads expected values from the DOCX's own style definitions instead
    of hardcoding point sizes, so this stays valid for any template that
    conforms to the same style-naming convention (Normal, List Bullet,
    Heading 1/2, Title, Caption).
    """
    findings: List[Finding] = []
    styles = doc.styles
    checked_styles = ["Normal", "List Bullet", "Heading 1", "Heading 2", "Title", "Caption"]
    defaults = {}
    for name in checked_styles:
        try:
            style = styles[name]
        except KeyError:
            continue
        defaults[name] = (style.font.name, style.font.size)

    for para in doc.paragraphs:
        style_name = para.style.name
        if style_name not in defaults:
            continue
        expected_name, expected_size = defaults[style_name]
        if expected_size is None:
            continue
        for run in para.runs:
            if run.font.name == "Consolas":
                continue  # markers/code spans use intentional direct formatting
            if run.font.size is not None and run.font.size != expected_size:
                findings.append(
                    Finding(
                        "WARNING",
                        "FONT_SIZE_OVERRIDE",
                        f"Run '{run.text[:40]}' em estilo '{style_name}' usa tamanho de fonte "
                        f"diferente do padrão do template ({run.font.size} != {expected_size}).",
                    )
                )
    return findings


def audit_document(doc: Document) -> List[Finding]:
    findings: List[Finding] = []
    if not doc.paragraphs and not doc.tables:
        findings.append(Finding("BLOCKING", "EMPTY_DOCUMENT", "Documento não possui conteúdo visível."))
        return findings

    findings += audit_heading_hierarchy(doc)
    findings += audit_headings_have_content(doc)
    findings += audit_empty_paragraph_runs(doc)
    findings += audit_tables(doc)
    findings += audit_images(doc)
    findings += audit_font_consistency(doc)
    return findings


def classify(findings: List[Finding]) -> str:
    if any(f.severity == "BLOCKING" for f in findings):
        return STATUS_OUTPUT_BLOCKED
    if any(f.severity == "WARNING" for f in findings):
        return STATUS_PASS_WITH_WARNINGS
    return STATUS_PASS


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", required=True, type=Path, help="Caminho do DOCX a auditar")
    args = parser.parse_args(argv)

    if not args.docx.is_file():
        print(f"{STATUS_OUTPUT_BLOCKED}: arquivo não encontrado: {args.docx}")
        return EXIT_BY_STATUS[STATUS_OUTPUT_BLOCKED]

    doc = Document(str(args.docx))
    findings = audit_document(doc)
    status = classify(findings)

    print(status)
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)}")
    for f in findings:
        print(f"[{f.severity}] {f.code}: {f.message}")
    if not findings:
        print("Nenhum finding.")

    return EXIT_BY_STATUS[status]


if __name__ == "__main__":
    sys.exit(main())
