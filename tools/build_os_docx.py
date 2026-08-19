#!/usr/bin/env python3
"""Generic DOCX materializer for `.osFactory` Ordens de Serviço (OS).

Converts the canonical OS Markdown (`05-output/<demanda>/OS-<demanda>.md`)
into a DOCX document, using `04-templates/docx/os-padrao.docx` as the
VISUAL source of truth (Tria identity: colors, fonts, header/footer,
cover page) while the Markdown remains the CONTENT source of truth
(functional specification, escopo, esforço, condições comerciais,
aceite).

This script is demand-agnostic by design: it never hardcodes a client
name, project code, or demand identifier. Everything demand-specific
comes from the Markdown file, the CLI arguments, and two small local
sources of institutional/state data:

  - `config/os-factory.json` (versioned): institutional defaults
    (Contratante, Executor, Autor, formato do código, método de
    aceite). See os-rules.md, OS-CODE-002.
  - `01-analysis/_runtime/os-registry.json` (git-ignored, see
    `tools/os_registry.py`): demand -> código da OS, so the same
    demand always keeps the same `OS-<AAAA>-<NNNN>` across
    regenerations. See os-rules.md, OS-CODE-001.

Gate discipline (see os-rules.md, OS-QA-003 / OS-QA-004):
  - FUNCTIONAL_BLOCKED is a content/specification problem and belongs to
    `os-validator-agent` / `validation.md`, evaluated BEFORE this script
    is ever invoked. This script does not re-judge functional quality.
    If `--validation-status BLOCKED` is passed explicitly, this script
    refuses to produce a DOCX.
  - OUTPUT_BLOCKED is a materialization failure (missing template,
    missing figure file, DOCX write error, etc.).
  - Acceptance readiness (READY_FOR_ACCEPTANCE / NOT_READY_FOR_ACCEPTANCE)
    is a SEPARATE, documental/commercial assessment (missing OS code,
    missing valor, missing forma de pagamento, etc.). It never turns a
    valid functional PASS/PASS_WITH_WARNINGS into FUNCTIONAL_BLOCKED,
    and it never blocks DOCX generation — it only affects the
    "Status documental" shown on the cover (Em elaboração / Para
    aceite) and is reported back to the caller.
  - "Aprovada" is NEVER inferred: it is only set via explicit
    `--approved` (+ `--approved-by`), never derived automatically.

Usage:
    python tools/build_os_docx.py \\
        --demand <demanda> \\
        --markdown 05-output/<demanda>/OS-<demanda>.md \\
        --template 04-templates/docx/os-padrao.docx \\
        --output 05-output/<demanda>/OS-<demanda>.docx \\
        --validation-status PASS_WITH_WARNINGS
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parent))
import os_registry  # noqa: E402
import os_commercial  # noqa: E402

# --------------------------------------------------------------------------
# Exit / status vocabulary
# --------------------------------------------------------------------------

STATUS_OK = "OK"
STATUS_FUNCTIONAL_BLOCKED = "FUNCTIONAL_BLOCKED"
STATUS_OUTPUT_BLOCKED = "OUTPUT_BLOCKED"

EXIT_OK = 0
EXIT_FUNCTIONAL_BLOCKED = 3
EXIT_OUTPUT_BLOCKED = 2

# Safety margin subtracted from the template's own computed usable width
# (page width minus left/right margins), so a figure never touches the
# page edges even after Word's own rendering rounding.
USABLE_WIDTH_SAFETY_INCHES = 0.15
EMU_PER_INCH = 914400

REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_CONFIG_PATH = REPO_ROOT / "config" / "os-factory.json"
DEFAULT_REGISTRY_PATH = os_registry.DEFAULT_REGISTRY_PATH
DEFAULT_COMMERCIAL_POLICY_PATH = os_commercial.DEFAULT_POLICY_PATH
DEFAULT_COMMERCIAL_REGISTRY_PATH = os_commercial.DEFAULT_COMMERCIAL_REGISTRY_PATH

MARKER_KEYWORDS = (
    "OPEN_QUESTION",
    "DISCOVERY_ITEM",
    "INFERENCE",
    "FUNCTIONAL_CONFLICT",
    "ARCHITECTURAL_CONFLICT",
    "DOCUMENTAL_CONFLICT",
    "CONFLICT_UNCLASSIFIED",
)
# Matches "[KEYWORD: ...]" or "[KEYWORD D1: ...]" (item ids like "D1"),
# with or without the surrounding brackets already stripped by the
# backtick/code-span tokenizer.
MARKER_PATTERN = re.compile(
    r"^\[?(" + "|".join(MARKER_KEYWORDS) + r")\b[^:]*:"
)

FIGURE_PATTERN = re.compile(
    r"^`?\[FIGURA:\s*(?P<description>.*?)\s*\|\s*fonte:\s*(?P<source>.*?)\]`?$"
)

# Cover-page metadata labels, in the order the official template lays
# them out (see 04-templates/docx/os-padrao.docx). Only used to detect
# whether a given template actually ships a compatible cover table —
# the generator degrades gracefully (skips cover-filling) against any
# template that doesn't define one, e.g. a minimal test fixture.
COVER_LABELS = [
    "código da os",
    "nome da demanda",
    "contratante",
    "cliente/projeto",
    "executor",
    "data de emissão",
    "versão",
    "status",
]


class MaterializationError(Exception):
    """Raised for OUTPUT_BLOCKED conditions (materialization, not content)."""


class FunctionalGateError(Exception):
    """Raised when the caller explicitly signals a FUNCTIONAL_BLOCKED gate."""


# --------------------------------------------------------------------------
# Markdown parsing (structural, content-preserving — never rewrites text)
# --------------------------------------------------------------------------

@dataclass
class Block:
    kind: str  # heading | paragraph | bullet | table | figure | quote
    level: int = 0
    text: str = ""
    rows: List[List[str]] = field(default_factory=list)
    description: str = ""
    source: str = ""


def _split_table_row(line: str) -> List[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _is_table_separator(cells: List[str]) -> bool:
    return all(re.fullmatch(r":?-{2,}:?", c.strip()) is not None for c in cells if c.strip())


def parse_markdown(text: str) -> List[Block]:
    """Parse the OS Markdown into a flat list of structural blocks.

    Intentionally simple and line-oriented (see design note in the repo
    conversation / commit context): the `.osFactory` documenter agent
    produces one logical block per non-blank line, so no soft-wrap
    joining is required. This keeps the parser predictable and avoids
    any risk of silently merging or rewording adjacent sentences.
    """
    lines = text.splitlines()
    blocks: List[Block] = []
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # HTML comments (template instructions) — never materialized.
        if stripped.startswith("<!--"):
            while i < n and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            blocks.append(Block(kind="heading", level=len(m.group(1)), text=m.group(2).strip()))
            i += 1
            continue

        # Figure marker (own line)
        fm = FIGURE_PATTERN.match(stripped)
        if fm:
            blocks.append(
                Block(
                    kind="figure",
                    description=fm.group("description").strip(),
                    source=fm.group("source").strip(),
                )
            )
            i += 1
            continue

        # Table block
        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [_split_table_row(l) for l in table_lines]
            if len(rows) >= 2 and _is_table_separator(rows[1]):
                rows = [rows[0]] + rows[2:]
            blocks.append(Block(kind="table", rows=rows))
            continue

        # Blockquote (exact system messages)
        if stripped.startswith(">"):
            blocks.append(Block(kind="quote", text=stripped.lstrip(">").strip()))
            i += 1
            continue

        # Bullet
        if stripped.startswith("- "):
            blocks.append(Block(kind="bullet", text=stripped[2:].strip()))
            i += 1
            continue

        # Plain paragraph (one line = one block, per the design note above)
        blocks.append(Block(kind="paragraph", text=stripped))
        i += 1

    return blocks


# --------------------------------------------------------------------------
# Inline formatting: bold (**...**) and code spans (`...`), incl. markers
# --------------------------------------------------------------------------

_INLINE_TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def iter_inline_runs(text: str):
    """Yield (text, is_bold, is_code) tuples for a line of inline markdown.

    Content is never altered — only classified for run formatting. Code
    spans are used verbatim for `[OPEN_QUESTION: ...]`,
    `[DISCOVERY_ITEM: ...]`, `[INFERENCE: ...]` and `CONFLICT` subtype
    markers, so they render distinctly without any special-casing beyond
    recognizing they are backtick-quoted, exactly like any other inline
    code span in the source Markdown.
    """
    for token in _INLINE_TOKEN.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) >= 4:
            yield token[2:-2], True, False
        elif token.startswith("`") and token.endswith("`") and len(token) >= 2:
            yield token[1:-1], False, True
        else:
            yield token, False, False


def _is_marker_text(inner: str) -> bool:
    return bool(MARKER_PATTERN.match(inner.strip()))


# Direct run formatting (not named character styles) is used for code
# spans and markers, on purpose: it makes the generator work correctly
# against ANY conformant template, even one that does not define the
# optional "OS Marker"/"OS Code" character styles the official
# `os-padrao.docx` ships with. Named paragraph styles (Normal,
# Heading 1/2, List Bullet, Title, Caption) are still used, since those
# are virtually guaranteed to exist in any Word-compatible template.
_MARKER_COLOR = RGBColor(0x7A, 0x1F, 0x1F)
_CODE_COLOR = RGBColor(0x33, 0x33, 0x33)
_MONOSPACE_FONT = "Consolas"


# Monospace fonts commonly render visually larger than proportional body
# text at the same nominal point size (metrics differ, and font
# substitution on non-Windows renderers makes this worse). Code/marker
# runs are sized explicitly, slightly under body size, so they stay
# visually harmonious regardless of which monospace font is actually
# available on the rendering machine.
_CODE_SIZE = Pt(10)


def add_inline_runs(paragraph, text: str) -> None:
    for chunk, is_bold, is_code in iter_inline_runs(text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if is_bold:
            run.bold = True
        if is_code:
            run.font.name = _MONOSPACE_FONT
            run.font.size = _CODE_SIZE
            if _is_marker_text(chunk):
                run.italic = True
                run.font.color.rgb = _MARKER_COLOR
            else:
                run.font.color.rgb = _CODE_COLOR


# --------------------------------------------------------------------------
# Document metadata extraction (Identificação / Controle de versão /
# Condições Comerciais / Aceite) — read-only, never rewrites the
# Markdown. Used to fill the DOCX cover page and to compute acceptance
# readiness. Field labels are matched case-insensitively and tolerant of
# minor punctuation/spacing variation, but the underlying VALUES are
# always used verbatim (never reworded).
# --------------------------------------------------------------------------

def _normalize_label(label: str) -> str:
    norm = label.strip().rstrip(":").strip().lower()
    norm = re.sub(r"\s*/\s*", "/", norm)
    norm = re.sub(r"\s+", " ", norm)
    return norm


def _normalize_heading(text: str) -> str:
    # Strip a leading "N. " / "N.N " numbering prefix, if present.
    stripped = re.sub(r"^\d+(\.\d+)*\.?\s*", "", text.strip())
    return stripped.strip().lower()


def _section_blocks(blocks: List[Block], heading_match) -> List[Block]:
    """Return the blocks between the first heading matched by
    `heading_match(normalized_text) -> bool` and the next heading."""
    collecting = False
    out: List[Block] = []
    for b in blocks:
        if b.kind == "heading":
            if collecting:
                break
            if heading_match(_normalize_heading(b.text)):
                collecting = True
            continue
        if collecting:
            out.append(b)
    return out


def _bullet_fields(section: List[Block]) -> Dict[str, str]:
    fields: Dict[str, str] = {}
    for b in section:
        if b.kind != "bullet" or ":" not in b.text:
            continue
        label, value = b.text.split(":", 1)
        fields[_normalize_label(label)] = value.strip()
    return fields


def _last_table_row_dict(section: List[Block]) -> Dict[str, str]:
    for b in section:
        if b.kind == "table" and b.rows:
            header = [(_normalize_label(c)) for c in b.rows[0]]
            data_rows = b.rows[1:]
            if not data_rows:
                return {}
            last = data_rows[-1]
            return {header[i]: (last[i] if i < len(last) else "") for i in range(len(header))}
    return {}


def _field_is_absent(value: Optional[str]) -> bool:
    if value is None:
        return True
    stripped = value.strip()
    if not stripped:
        return True
    if "OPEN_QUESTION" in stripped:
        return True
    return False


def extract_document_metadata(blocks: List[Block]) -> Dict[str, object]:
    """Extract everything the cover page / acceptance-readiness logic
    needs from the Markdown, without ever rewriting it."""
    identificacao = _bullet_fields(
        _section_blocks(blocks, lambda h: h == "identificação")
    )
    controle_versao = _last_table_row_dict(
        _section_blocks(blocks, lambda h: h == "controle de versão")
    )
    condicoes = _bullet_fields(
        _section_blocks(blocks, lambda h: "condições comerciais" in h)
    )
    aceite_present = bool(_section_blocks(blocks, lambda h: h == "aceite"))

    return {
        "identificacao": identificacao,
        "controle_versao": controle_versao,
        "condicoes_comerciais": condicoes,
        "aceite_present": aceite_present,
    }


def _infer_year(demand: str, explicit_year: Optional[str], identificacao: Dict[str, str]) -> str:
    if explicit_year:
        if not re.fullmatch(r"\d{4}", explicit_year):
            raise MaterializationError(f"--year inválido: {explicit_year!r} (esperado AAAA)")
        return explicit_year

    m = re.match(r"^\d{2}-\d{2}-(\d{4})-", demand)
    if m:
        return m.group(1)

    data = identificacao.get("data", "")
    m2 = re.search(r"\b(\d{4})\b", data)
    if m2:
        return m2.group(1)

    raise MaterializationError(
        "não foi possível inferir o ano para gerar um novo código de OS "
        "(demanda não segue o padrão DD-MM-AAAA-<slug> e a Identificação "
        "não informa uma Data reconhecível) — informe --year explicitamente"
    )


def load_institutional_config(path: Path) -> Dict[str, str]:
    if not path.is_file():
        raise MaterializationError(f"configuração institucional não encontrada: {path}")
    with path.open("r", encoding="utf-8") as f:
        data = json.load(f)
    required = ("contratante_padrao", "executor_padrao", "autor_padrao", "metodo_padrao_aceite")
    missing = [k for k in required if k not in data]
    if missing:
        raise MaterializationError(f"config/os-factory.json incompleto — faltando: {missing}")
    return data


REQUIRED_COMMERCIAL_FIELDS = {
    "horas contratadas": "Horas contratadas",
    "valor total da os": "Valor da OS",
    "forma/condição de pagamento": "Forma de pagamento",
}


def compute_acceptance_readiness(
    *,
    functional_status: Optional[str],
    os_code: Optional[str],
    versao: Optional[str],
    contratante: Optional[str],
    executor: Optional[str],
    condicoes_comerciais: Dict[str, str],
) -> Tuple[str, List[str]]:
    """Documental/commercial readiness — SEPARATE from the functional gate.

    Never turns a valid functional PASS/PASS_WITH_WARNINGS into
    FUNCTIONAL_BLOCKED (see os-rules.md, OS-QA-004): it only reports
    what is still missing for the OS to be sent for acceptance.
    """
    missing: List[str] = []
    if functional_status == "BLOCKED":
        missing.append("Validação funcional (BLOCKED)")
    if not os_code:
        missing.append("Código da OS")
    if not versao:
        missing.append("Versão")
    if not contratante:
        missing.append("Contratante")
    if not executor:
        missing.append("Executor")
    for norm_label, human_label in REQUIRED_COMMERCIAL_FIELDS.items():
        if _field_is_absent(condicoes_comerciais.get(norm_label)):
            missing.append(human_label)

    readiness = "NOT_READY_FOR_ACCEPTANCE" if missing else "READY_FOR_ACCEPTANCE"
    return readiness, missing


# --------------------------------------------------------------------------
# DOCX rendering
# --------------------------------------------------------------------------

HEADING_STYLE_BY_LEVEL = {1: "Title", 2: "Heading 1", 3: "Heading 2"}


def _clear_placeholder(doc: Document) -> None:
    """Remove the template's placeholder body paragraph, if present.

    Searches every paragraph (not just index 0) because the official
    template now places a cover page — title, subtitle, metadata table
    — before the placeholder paragraph. Only removes a paragraph that
    matches the known placeholder text; never deletes real content.
    """
    for p in list(doc.paragraphs):
        if "template oficial da .osfactory" in p.text.strip().lower():
            element = p._element
            element.getparent().remove(element)
            return


def _looks_like_cover_table(table) -> bool:
    """Heuristic: only treat a table as the cover metadata table when
    most of its first-column labels match the official template's
    known cover labels. Prevents accidentally corrupting a generic/test
    template's own body table."""
    if not table.rows:
        return False
    labels = {_normalize_label(row.cells[0].text) for row in table.rows}
    matches = sum(1 for expected in COVER_LABELS if expected in labels)
    return matches >= max(4, len(COVER_LABELS) // 2)


def _set_value_cell(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    add_inline_runs(p, text) if text else None


def _fill_cover_table(cover_table, cover_values: Dict[str, str]) -> None:
    for row in cover_table.rows:
        label = _normalize_label(row.cells[0].text)
        if label in cover_values:
            _set_value_cell(row.cells[1], cover_values[label])

    approval = cover_values.get("__approval_extra__")
    if approval:
        for extra_label, extra_value in approval:
            new_row = cover_table.add_row()
            label_p = new_row.cells[0].paragraphs[0]
            label_run = label_p.add_run(f"{extra_label}:")
            label_run.bold = True
            _set_value_cell(new_row.cells[1], extra_value)


def _override_autor_column(rows: List[List[str]], autor_padrao: str) -> List[List[str]]:
    """Force the 'Autor' column of the Controle de versão table to the
    institutional default (see os-rules.md, OS-CODE-002) — this is
    document-control metadata, not functional content, so overriding it
    at materialization time (rather than trusting whatever the Markdown
    happened to record) is a deliberate, documented exception to
    "never rewrite content"."""
    if not rows:
        return rows
    header = [_normalize_label(c) for c in rows[0]]
    if "autor" not in header:
        return rows
    idx = header.index("autor")
    new_rows = [rows[0]]
    for r in rows[1:]:
        r = list(r)
        if idx < len(r):
            r[idx] = autor_padrao
        new_rows.append(r)
    return new_rows


def _render_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.paragraphs[0].text = ""
            para = cell.paragraphs[0]
            add_inline_runs(para, cell_text)
            for run in para.runs:
                run.font.size = Pt(10)
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True
    # blank line after table for breathing room
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def _resolve_figure(source: str, figures_dir: Optional[Path]) -> Path:
    if figures_dir is None:
        raise MaterializationError(
            f"figura referenciada ({source}) mas nenhum diretório de figuras foi informado"
        )
    candidate = figures_dir / source
    if not candidate.is_file():
        raise MaterializationError(
            f"figura ausente: esperado em {candidate} (referenciada como '{source}')"
        )
    return candidate


def _usable_width_inches(doc: Document) -> float:
    """Compute the usable page width from the TEMPLATE's own section
    margins (never hardcoded), so figures stay correctly sized no matter
    what page size/margins a given template defines."""
    section = doc.sections[0]
    usable_emu = section.page_width - section.left_margin - section.right_margin
    return max(usable_emu / EMU_PER_INCH - USABLE_WIDTH_SAFETY_INCHES, 1.0)


def _render_figure(doc: Document, block: Block, figures_dir: Optional[Path]) -> None:
    image_path = _resolve_figure(block.source, figures_dir)

    # Validate the file is actually a readable image before handing it to
    # python-docx (which would otherwise raise a less legible error).
    with Image.open(image_path) as img:
        img.verify()

    # Only width is set — python-docx preserves the image's native aspect
    # ratio automatically, so proportion is never distorted, and the
    # figure never exceeds the page's usable width.
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(_usable_width_inches(doc)))

    caption = doc.add_paragraph(style="Caption")
    caption.text = f"Figura — {block.description} (fonte: {block.source})"


_FIXED_ZIP_DATETIME = (1980, 1, 1, 0, 0, 0)


def _normalize_zip_timestamps(docx_path: Path) -> None:
    """Rewrite every ZIP entry's timestamp to a fixed value.

    A .docx file is a ZIP container; python-docx (via the stdlib
    zipfile module) stamps each internal entry with the current
    wall-clock time on save. That alone would make two runs over the
    same Markdown/template produce different bytes even though the
    visible content is identical. Normalizing the timestamps makes the
    output byte-for-byte deterministic.
    """
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(
        tmp_path, "w", zipfile.ZIP_DEFLATED
    ) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            item.date_time = _FIXED_ZIP_DATETIME
            dst.writestr(item, data)
    tmp_path.replace(docx_path)


def _apply_commercial_override(text: str, resolved: Dict[str, Optional[str]]) -> str:
    """Replace the value half of a 'Label: value' bullet with the
    hierarchy-resolved value (see tools/os_commercial.py), when the
    label matches a known commercial field and a resolved value exists.
    Preserves the bullet's own label text verbatim; never touches
    bullets outside CONDIÇÕES COMERCIAIS or labels it doesn't recognize."""
    if ":" not in text:
        return text
    label, _, _rest = text.partition(":")
    key = _normalize_label(label)
    if key in resolved and resolved[key]:
        return f"{label.strip()}: {resolved[key]}"
    return text


def render_docx(
    blocks: List[Block],
    template_path: Path,
    output_path: Path,
    figures_dir: Optional[Path],
    cover_values: Optional[Dict[str, str]],
    autor_padrao: Optional[str],
    resolved_commercial: Optional[Dict[str, Optional[str]]] = None,
) -> None:
    if not template_path.is_file():
        raise MaterializationError(f"template DOCX não encontrado: {template_path}")

    # Work on a temp copy — the original template is never opened for
    # writing and is never modified.
    with tempfile.TemporaryDirectory() as tmp:
        working_copy = Path(tmp) / "working.docx"
        shutil.copyfile(template_path, working_copy)

        doc = Document(str(working_copy))

        # Capture (and fill) the cover metadata table BEFORE any body
        # content is appended — doc.tables[0] at this point can only be
        # the template's own cover table, if it has one.
        if cover_values is not None and doc.tables and _looks_like_cover_table(doc.tables[0]):
            _fill_cover_table(doc.tables[0], cover_values)

        _clear_placeholder(doc)

        pending_heading = ""
        for block in blocks:
            if block.kind == "heading":
                pending_heading = _normalize_heading(block.text)
                style = HEADING_STYLE_BY_LEVEL.get(block.level, "Heading 2")
                p = doc.add_paragraph(style=style)
                add_inline_runs(p, block.text)
            elif block.kind == "paragraph":
                p = doc.add_paragraph(style="Normal")
                add_inline_runs(p, block.text)
            elif block.kind == "bullet":
                text = block.text
                if resolved_commercial and "condições comerciais" in pending_heading:
                    text = _apply_commercial_override(text, resolved_commercial)
                p = doc.add_paragraph(style="List Bullet")
                add_inline_runs(p, text)
            elif block.kind == "quote":
                p = doc.add_paragraph(style="Intense Quote")
                add_inline_runs(p, block.text)
            elif block.kind == "table":
                rows = block.rows
                if pending_heading == "controle de versão" and autor_padrao:
                    rows = _override_autor_column(rows, autor_padrao)
                _render_table(doc, rows)
            elif block.kind == "figure":
                _render_figure(doc, block, figures_dir)
            else:  # pragma: no cover - defensive
                raise MaterializationError(f"tipo de bloco desconhecido: {block.kind}")

        # Deterministic core properties: no wall-clock timestamps, so the
        # same Markdown + template always produce the same content.
        cp = doc.core_properties
        cp.author = ".osFactory"
        cp.last_modified_by = ".osFactory"
        cp.revision = 1

        output_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            doc.save(str(output_path))
            _normalize_zip_timestamps(output_path)
        except MaterializationError:
            raise
        except Exception as exc:  # pragma: no cover - defensive
            raise MaterializationError(f"falha ao salvar DOCX de saída: {exc}") from exc


# --------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------

def build(
    demand: str,
    markdown_path: Path,
    template_path: Path,
    output_path: Path,
    figures_dir: Optional[Path],
    validation_status: Optional[str],
    config_path: Path,
    registry_path: Path,
    explicit_os_code: Optional[str],
    explicit_year: Optional[str],
    approved: bool,
    approved_by: Optional[str],
    approved_date: Optional[str],
    acceptance_reference: Optional[str],
    commercial_policy_path: Path = DEFAULT_COMMERCIAL_POLICY_PATH,
    commercial_registry_path: Path = DEFAULT_COMMERCIAL_REGISTRY_PATH,
    explicit_horas: Optional[str] = None,
    explicit_valor: Optional[str] = None,
    explicit_pagamento_splits: Optional[List[Dict]] = None,
    explicit_commercial_note: Optional[str] = None,
) -> Dict[str, object]:
    if validation_status is not None and validation_status.upper() == "BLOCKED":
        raise FunctionalGateError(
            "validação funcional retornou BLOCKED — DOCX não pode ser gerado como "
            "documento final aprovado (ver os-rules.md, OS-QA-003). Corrija o "
            "conteúdo e revalide antes de materializar."
        )

    if not markdown_path.is_file():
        raise MaterializationError(f"Markdown de origem não encontrado: {markdown_path}")

    if approved and not approved_by:
        raise MaterializationError(
            "--approved exige --approved-by (o aceite nunca é inferido — ver os-rules.md, OS-CODE-003)"
        )

    text = markdown_path.read_text(encoding="utf-8")
    blocks = parse_markdown(text)
    meta = extract_document_metadata(blocks)
    identificacao = meta["identificacao"]
    controle_versao = meta["controle_versao"]
    condicoes = meta["condicoes_comerciais"]

    config = load_institutional_config(config_path)

    year = _infer_year(demand, explicit_year, identificacao)
    os_code = os_registry.resolve_os_code(
        demand, year, explicit_code=explicit_os_code, registry_path=registry_path
    )

    contratante = config["contratante_padrao"]
    executor = config["executor_padrao"]
    autor_padrao = config["autor_padrao"]
    metodo_aceite = config["metodo_padrao_aceite"]

    nome_demanda = identificacao.get("nome da demanda", "")
    versao = controle_versao.get("versão", "")
    data_emissao = controle_versao.get("data", "") or identificacao.get("data", "")

    cliente_projeto_parts = [
        p for p in (identificacao.get("solicitante"), identificacao.get("projeto")) if p
    ]
    cliente_projeto = " — ".join(cliente_projeto_parts)

    commercial_policy = os_commercial.load_policy(commercial_policy_path)
    resolved_commercial = os_commercial.resolve_commercial_terms(
        demand=demand,
        explicit_horas=explicit_horas,
        explicit_valor=explicit_valor,
        explicit_pagamento_splits=explicit_pagamento_splits,
        explicit_note=explicit_commercial_note,
        insumos=condicoes,
        policy=commercial_policy,
        registry_path=commercial_registry_path,
    )
    effective_condicoes = {k: v.value for k, v in resolved_commercial.items()}

    readiness, missing = compute_acceptance_readiness(
        functional_status=validation_status,
        os_code=os_code,
        versao=versao,
        contratante=contratante,
        executor=executor,
        condicoes_comerciais=effective_condicoes,
    )
    payment_field = resolved_commercial["forma/condição de pagamento"]
    if payment_field.source == os_commercial.PAYMENT_BY_MILESTONES_RECOMMENDED and "Forma de pagamento" in missing:
        missing = [
            "Forma de pagamento (PAYMENT_BY_MILESTONES_RECOMMENDED — definir distribuição por marcos)"
            if m == "Forma de pagamento" else m
            for m in missing
        ]

    if approved:
        document_status = "Aprovada"
    elif readiness == "READY_FOR_ACCEPTANCE":
        document_status = "Para aceite"
    else:
        document_status = "Em elaboração"

    cover_values: Dict[str, str] = {
        "código da os": os_code,
        "nome da demanda": nome_demanda,
        "contratante": contratante,
        "cliente/projeto": cliente_projeto,
        "executor": executor,
        "data de emissão": data_emissao,
        "versão": versao,
        "status": document_status,
    }
    if approved:
        extra = []
        if approved_date:
            extra.append(("Data do aceite", approved_date))
        extra.append(("Aprovado por", approved_by))
        if acceptance_reference:
            extra.append(("Referência do aceite", acceptance_reference))
        cover_values["__approval_extra__"] = extra

    render_docx(blocks, template_path, output_path, figures_dir, cover_values, autor_padrao, effective_condicoes)

    return {
        "os_code": os_code,
        "functional_validation": validation_status or "(não informado)",
        "acceptance_readiness": readiness,
        "missing": missing,
        "document_status": document_status,
        "metodo_aceite": metodo_aceite,
        "aceite_present_in_markdown": meta["aceite_present"],
        "output_path": str(output_path),
        "commercial": {
            k: {"value": v.value, "source": v.source} for k, v in resolved_commercial.items()
        },
    }


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--demand", required=True, help="Identificador da demanda")
    parser.add_argument("--markdown", required=True, type=Path, help="Caminho do OS-<demanda>.md")
    parser.add_argument("--template", required=True, type=Path, help="Caminho do template DOCX oficial")
    parser.add_argument("--output", required=True, type=Path, help="Caminho do DOCX de saída")
    parser.add_argument(
        "--figures-dir",
        type=Path,
        default=None,
        help="Diretório onde localizar arquivos de figura referenciados por [FIGURA: ... | fonte: <arquivo>]. "
        "Padrão: 00-inbox/<demand>/ relativo à raiz do repositório (inferida a partir deste script).",
    )
    parser.add_argument(
        "--validation-status",
        default=None,
        choices=["PASS", "PASS_WITH_WARNINGS", "BLOCKED"],
        help="Resultado do os-validator-agent para esta demanda. Se BLOCKED, o script recusa gerar o DOCX.",
    )
    parser.add_argument("--config", type=Path, default=DEFAULT_CONFIG_PATH, help="config/os-factory.json")
    parser.add_argument("--registry", type=Path, default=DEFAULT_REGISTRY_PATH, help="Registry local de código de OS")
    parser.add_argument("--os-code", default=None, help="Código da OS explícito/confirmado (OS-AAAA-NNNN)")
    parser.add_argument("--year", default=None, help="Ano para geração de novo código, se não inferível da demanda/Markdown")
    parser.add_argument("--approved", action="store_true", help="Marca Status documental como Aprovada (nunca inferido)")
    parser.add_argument("--approved-by", default=None, help="Obrigatório se --approved")
    parser.add_argument("--approved-date", default=None)
    parser.add_argument("--acceptance-reference", default=None)
    parser.add_argument("--commercial-policy", type=Path, default=DEFAULT_COMMERCIAL_POLICY_PATH, help="config/os-commercial-policy.json")
    parser.add_argument("--commercial-registry", type=Path, default=DEFAULT_COMMERCIAL_REGISTRY_PATH, help="Registry local de condições comerciais confirmadas")
    parser.add_argument("--horas-contratadas", default=None, help="Horas contratadas confirmadas explicitamente pelo usuário (maior autoridade — ver OS-COMMERCIAL-001)")
    parser.add_argument("--valor-total-os", default=None, help="Valor total da OS confirmado explicitamente (ex.: 'R$ 8.000,00')")
    parser.add_argument(
        "--forma-pagamento-marcos",
        default=None,
        help="JSON confirmado explicitamente: lista de {\"marco\": str, \"percentual\": number} somando 100. "
        "Os valores de cada parcela são sempre calculados a partir de --valor-total-os, nunca informados manualmente.",
    )
    parser.add_argument("--commercial-note", default=None, help="Nota de rastreabilidade interna (ex.: motivo de uma condição substituir um insumo anterior)")
    args = parser.parse_args(argv)

    pagamento_splits = None
    if args.forma_pagamento_marcos:
        try:
            pagamento_splits = json.loads(args.forma_pagamento_marcos)
        except json.JSONDecodeError as exc:
            print(f"{STATUS_OUTPUT_BLOCKED}: --forma-pagamento-marcos não é um JSON válido: {exc}")
            return EXIT_OUTPUT_BLOCKED

    figures_dir = args.figures_dir
    if figures_dir is None:
        default_dir = REPO_ROOT / "00-inbox" / args.demand
        figures_dir = default_dir if default_dir.is_dir() else None

    try:
        result = build(
            demand=args.demand,
            markdown_path=args.markdown,
            template_path=args.template,
            output_path=args.output,
            figures_dir=figures_dir,
            validation_status=args.validation_status,
            config_path=args.config,
            registry_path=args.registry,
            explicit_os_code=args.os_code,
            explicit_year=args.year,
            approved=args.approved,
            approved_by=args.approved_by,
            approved_date=args.approved_date,
            acceptance_reference=args.acceptance_reference,
            commercial_policy_path=args.commercial_policy,
            commercial_registry_path=args.commercial_registry,
            explicit_horas=args.horas_contratadas,
            explicit_valor=args.valor_total_os,
            explicit_pagamento_splits=pagamento_splits,
            explicit_commercial_note=args.commercial_note,
        )
    except FunctionalGateError as exc:
        print(f"{STATUS_FUNCTIONAL_BLOCKED}: {exc}")
        return EXIT_FUNCTIONAL_BLOCKED
    except (MaterializationError, os_commercial.CommercialError) as exc:
        print(f"{STATUS_OUTPUT_BLOCKED}: {exc}")
        return EXIT_OUTPUT_BLOCKED

    print(f"{STATUS_OK}: DOCX gerado em {result['output_path']}")
    print(f"Código da OS: {result['os_code']}")
    print(f"Functional validation: {result['functional_validation']}")
    print(f"Acceptance readiness: {result['acceptance_readiness']}")
    print(f"Document status: {result['document_status']}")
    print(f"Método de aceite: {result['metodo_aceite']}")
    for key, label in (
        ("horas contratadas", "Horas contratadas"),
        ("valor total da os", "Valor total da OS"),
        ("forma/condição de pagamento", "Forma de pagamento"),
    ):
        field = result["commercial"][key]
        print(f"{label}: {field['value'] if field['value'] else '(ausente)'} [{field['source']}]")
    if result["missing"]:
        print("Pendências para aceite:")
        for m in result["missing"]:
            print(f"- {m}")
    return EXIT_OK


if __name__ == "__main__":
    sys.exit(main())
